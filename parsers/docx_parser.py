# parsers/docx_parser.py
"""
DOCX resume parser using python-docx.
Extracts text from paragraphs and tables.
"""

from docx import Document
from pathlib import Path
from typing import Union

from utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(docx_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX file.

    Args:
        docx_path: Path to the DOCX file.

    Returns:
        Extracted text as a single string with elements separated by newlines.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file is not a .docx.
        Exception: If document reading fails.
    """
    docx_path = Path(docx_path)

    # Validate file
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    if docx_path.suffix.lower() != '.docx':
        raise ValueError(f"File is not a DOCX: {docx_path}")

    try:
        logger.info(f"Extracting text from {docx_path}")
        doc = Document(docx_path)

        text_parts = []

        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)

        if not text_parts:
            logger.warning(f"No text found in {docx_path}")

        logger.info(f"Extracted {len(text_parts)} text segments from {docx_path}")
        return "\n".join(text_parts)

    except Exception as e:
        logger.error(f"Failed to extract text from {docx_path}: {e}")
        raise